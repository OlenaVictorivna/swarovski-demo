"""
Swarovski · AI-Readiness Scoring Intelligence Platform
Luxury corporate UI · Persistent results · Conversational Q&A
"""

import streamlit as st
import anthropic
import json
import re
import io
import os
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────── PAGE CONFIG ────────────────────────────
st.set_page_config(
    page_title="Swarovski · AI Intelligence Platform",
    page_icon="💎",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────── CONSTANTS ──────────────────────────────
PROMPT_FILE = Path(__file__).parent / "scoring_prompt.docx"

QUADRANT_EMOJI = {
    "FAST TRACK":         "⚡",
    "STRATEGIC PRIORITY": "🎯",
    "OPTIMISE FIRST":     "🔧",
    "DEPRIORITISE":       "🔴",
    "NOT A PROCEDURE":    "❌",
}
QUADRANT_HEX = {
    "FAST TRACK":         "00C896",
    "STRATEGIC PRIORITY": "3D8FD1",
    "OPTIMISE FIRST":     "F0A500",
    "DEPRIORITISE":       "E05555",
    "NOT A PROCEDURE":    "8899AA",
}
QUADRANT_LABEL = {
    "FAST TRACK":         "High Value · Low Complexity",
    "STRATEGIC PRIORITY": "High Value · High Complexity",
    "OPTIMISE FIRST":     "Low Value · Low Complexity",
    "DEPRIORITISE":       "Low Value · High Complexity",
}
SCORE_BG  = {3: "00C896", 2: "F0A500", 1: "E05555"}
HEADER_BG = "0D1B3E"
LABEL_BG  = "EEF2FF"


# ─────────────────────────── CSS ────────────────────────────────────
CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;600;700&family=DM+Sans:wght@300;400;500;600&display=swap');

/* ── Root vars ── */
:root {
    --navy:      #0D1B3E;
    --navy-mid:  #152550;
    --navy-lite: #1E3469;
    --gold:      #C9A84C;
    --gold-lite: #E8D08A;
    --crystal:   #3D8FD1;
    --green:     #00C896;
    --amber:     #F0A500;
    --red:       #E05555;
    --white:     #FFFFFF;
    --off-white: #F5F7FF;
    --text-main: #1A2B52;
    --text-sub:  #4A5980;
    --border:    rgba(13,27,62,0.10);
    --shadow:    0 4px 24px rgba(13,27,62,0.10);
}

/* ── App shell ── */
.stApp { background: var(--off-white); font-family: 'DM Sans', sans-serif; }
.main  .block-container { padding: 2rem 2.5rem 3rem; max-width: 1100px; }

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, var(--navy) 0%, var(--navy-mid) 60%, #0A1628 100%);
    border-right: 1px solid rgba(201,168,76,0.25);
}
[data-testid="stSidebar"] * { color: #CDD8F0 !important; font-family: 'DM Sans', sans-serif; }
[data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
    color: var(--gold-lite) !important;
    font-family: 'Cormorant Garamond', serif !important;
}
[data-testid="stSidebar"] .stFileUploader label { color: var(--gold-lite) !important; }
[data-testid="stSidebar"] hr { border-color: rgba(201,168,76,0.20) !important; }

/* ── Sidebar logo strip ── */
.sidebar-logo {
    text-align: center;
    padding: 1.5rem 1rem 1rem;
    border-bottom: 1px solid rgba(201,168,76,0.25);
    margin-bottom: 1.5rem;
}
.sidebar-logo .brand {
    font-family: 'Cormorant Garamond', serif;
    font-size: 1.4rem;
    font-weight: 700;
    color: #FFFFFF !important;
    letter-spacing: 0.18em;
    display: block;
}
.sidebar-logo .sub {
    font-size: 0.68rem;
    color: var(--gold) !important;
    letter-spacing: 0.22em;
    text-transform: uppercase;
    margin-top: 2px;
}
.diamond { font-size: 2.4rem; line-height: 1.2; display: block; color: #FFFFFF !important; }

/* ── Page title ── */
.page-header {
    padding: 2rem 0 1.5rem;
    border-bottom: 2px solid var(--border);
    margin-bottom: 2rem;
}
.page-title {
    font-family: 'Cormorant Garamond', serif;
    font-size: 2.1rem;
    font-weight: 700;
    color: var(--navy);
    letter-spacing: 0.04em;
    margin: 0;
}
.page-sub {
    font-size: 0.85rem;
    color: var(--text-sub);
    margin-top: 0.3rem;
    letter-spacing: 0.03em;
}
.gold-bar {
    width: 48px; height: 3px;
    background: linear-gradient(90deg, var(--gold), var(--gold-lite));
    border-radius: 2px;
    margin: 0.7rem 0;
}

/* ── Score cards ── */
.score-grid { display: flex; gap: 1rem; margin: 1.5rem 0; }
.score-card {
    flex: 1;
    border-radius: 14px;
    padding: 1.6rem 1.2rem;
    text-align: center;
    position: relative;
    overflow: hidden;
}
.score-card::before {
    content: '';
    position: absolute; top: 0; left: 0; right: 0; bottom: 0;
    background: linear-gradient(135deg, rgba(255,255,255,0.12) 0%, transparent 60%);
}
.score-card .label {
    font-size: 0.72rem;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: rgba(255,255,255,0.80);
    margin-bottom: 0.5rem;
    font-weight: 500;
}
.score-card .value {
    font-family: 'Cormorant Garamond', serif;
    font-size: 3.8rem;
    font-weight: 700;
    color: #FFFFFF;
    line-height: 1;
}
.score-card .unit {
    font-size: 0.72rem;
    color: rgba(255,255,255,0.65);
    margin-top: 0.3rem;
}

/* ── Section headers ── */
.section-title {
    font-family: 'Cormorant Garamond', serif;
    font-size: 1.2rem;
    font-weight: 600;
    color: var(--navy);
    letter-spacing: 0.06em;
    text-transform: uppercase;
    margin: 2rem 0 1rem;
    display: flex;
    align-items: center;
    gap: 0.6rem;
}
.section-title::after {
    content: '';
    flex: 1;
    height: 1px;
    background: var(--border);
}

/* ── Dimension rows ── */
.dim-row {
    display: flex;
    align-items: stretch;
    gap: 0;
    margin-bottom: 8px;
    border-radius: 10px;
    overflow: hidden;
    box-shadow: var(--shadow);
}
.dim-badge {
    width: 52px; min-width: 52px;
    display: flex; align-items: center; justify-content: center;
    font-family: 'Cormorant Garamond', serif;
    font-size: 1.8rem;
    font-weight: 700;
    color: #FFFFFF;
}
.dim-label {
    background: var(--off-white);
    padding: 0.7rem 1rem;
    min-width: 180px; max-width: 180px;
    display: flex; flex-direction: column; justify-content: center;
    border-right: 1px solid var(--border);
}
.dim-label .code { font-size: 0.7rem; letter-spacing: 0.1em; color: var(--text-sub); font-weight: 500; }
.dim-label .name { font-size: 0.82rem; font-weight: 600; color: var(--text-main); margin-top: 2px; }
.dim-label .tag  { font-size: 0.68rem; color: var(--text-sub); margin-top: 2px; }
.dim-rationale {
    background: #FFFFFF;
    padding: 0.7rem 1.1rem;
    flex: 1;
    font-size: 0.82rem;
    color: var(--text-main);
    display: flex; align-items: center;
    line-height: 1.5;
}

/* ── Quadrant banner ── */
.quadrant-banner {
    border-radius: 14px;
    padding: 1.8rem 2rem;
    display: flex;
    align-items: center;
    gap: 2rem;
    margin: 1.5rem 0;
    position: relative;
    overflow: hidden;
}
.quadrant-banner::before {
    content: '';
    position: absolute; top: 0; left: 0; right: 0; bottom: 0;
    background: linear-gradient(135deg, rgba(255,255,255,0.14) 0%, transparent 55%);
}
.q-emoji  { font-size: 3rem; line-height: 1; }
.q-title  { font-family: 'Cormorant Garamond', serif; font-size: 1.9rem; font-weight: 700; color: #FFFFFF; }
.q-sub    { font-size: 0.8rem; color: rgba(255,255,255,0.78); margin-top: 3px; letter-spacing: 0.08em; }
.q-scores { margin-left: auto; text-align: right; }
.q-scores span { display: block; font-size: 0.78rem; color: rgba(255,255,255,0.78); }
.q-scores strong { font-family: 'Cormorant Garamond', serif; font-size: 1.4rem; color: #FFFFFF; }

/* ── Exec summary card ── */
.exec-card {
    background: #FFFFFF;
    border: 1px solid var(--border);
    border-left: 4px solid var(--gold);
    border-radius: 10px;
    padding: 1.4rem 1.6rem;
    font-size: 0.88rem;
    line-height: 1.75;
    color: var(--text-main);
    box-shadow: var(--shadow);
    margin: 1rem 0 1.5rem;
}

/* ── Chat interface ── */
.chat-container {
    background: #FFFFFF;
    border: 1px solid var(--border);
    border-radius: 14px;
    overflow: hidden;
    box-shadow: var(--shadow);
    margin-top: 1rem;
}
.chat-header {
    background: linear-gradient(90deg, var(--navy) 0%, var(--navy-lite) 100%);
    padding: 1rem 1.5rem;
    display: flex;
    align-items: center;
    gap: 0.8rem;
}
.chat-header .ch-title {
    font-family: 'Cormorant Garamond', serif;
    font-size: 1.05rem;
    font-weight: 600;
    color: #FFFFFF;
    letter-spacing: 0.05em;
}
.chat-header .ch-sub { font-size: 0.72rem; color: rgba(255,255,255,0.65); }
.chat-dot { width: 8px; height: 8px; border-radius: 50%; background: var(--green); display: inline-block; }

.msg-user {
    background: linear-gradient(135deg, var(--navy) 0%, var(--navy-lite) 100%);
    color: #FFFFFF;
    border-radius: 16px 16px 4px 16px;
    padding: 0.75rem 1.1rem;
    margin: 0.4rem 0;
    max-width: 72%;
    margin-left: auto;
    font-size: 0.85rem;
    line-height: 1.55;
    box-shadow: 0 2px 10px rgba(13,27,62,0.18);
}
.msg-bot {
    background: var(--off-white);
    border: 1px solid var(--border);
    color: var(--text-main);
    border-radius: 16px 16px 16px 4px;
    padding: 0.75rem 1.1rem;
    margin: 0.4rem 0;
    max-width: 86%;
    font-size: 0.85rem;
    line-height: 1.65;
}
.msg-bot strong { color: var(--navy); }
.msg-avatar-bot {
    font-size: 1.1rem;
    margin-right: 6px;
    vertical-align: middle;
}

/* ── Metadata table ── */
.meta-table { width: 100%; border-collapse: collapse; margin: 0.5rem 0 1.5rem; }
.meta-table td { padding: 7px 12px; font-size: 0.82rem; border-bottom: 1px solid var(--border); }
.meta-table td:first-child {
    font-weight: 600; color: var(--navy); background: var(--off-white);
    width: 36%; letter-spacing: 0.02em;
}
.meta-table td:last-child { color: var(--text-main); background: #FFFFFF; }
.meta-table tr:last-child td { border-bottom: none; }

/* ── Download area ── */
.download-strip {
    background: linear-gradient(90deg, var(--navy) 0%, var(--navy-lite) 100%);
    border-radius: 12px;
    padding: 1.2rem 1.6rem;
    display: flex; align-items: center; justify-content: space-between;
    margin: 1.5rem 0;
}
.download-strip .ds-text { color: rgba(255,255,255,0.80); font-size: 0.82rem; }
.download-strip .ds-title { color: #FFFFFF; font-size: 0.95rem; font-weight: 600; }

/* ── How-it-works ── */
.hiw-card {
    background: #FFFFFF;
    border: 1px solid var(--border);
    border-radius: 12px;
    padding: 1.4rem;
    box-shadow: var(--shadow);
}
.step-item { display: flex; gap: 1rem; margin-bottom: 1rem; align-items: flex-start; }
.step-num {
    width: 28px; height: 28px; border-radius: 50%; min-width: 28px;
    background: linear-gradient(135deg, var(--navy), var(--navy-lite));
    color: #FFFFFF; font-size: 0.75rem; font-weight: 700;
    display: flex; align-items: center; justify-content: center;
}
.step-text { font-size: 0.83rem; color: var(--text-main); line-height: 1.55; padding-top: 4px; }
.step-text strong { color: var(--navy); }

/* ── Misc ── */
.pill {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 20px;
    font-size: 0.7rem;
    font-weight: 600;
    letter-spacing: 0.06em;
    text-transform: uppercase;
}
.pill-green  { background: rgba(0,200,150,0.15); color: #007A5E; }
.pill-amber  { background: rgba(240,165,0,0.15);  color: #8A5C00; }
.pill-red    { background: rgba(224,85,85,0.15);   color: #9C2424; }
.pill-blue   { background: rgba(61,143,209,0.15);  color: #1A5A8C; }

/* ── Override Streamlit defaults ── */
.stButton > button {
    font-family: 'DM Sans', sans-serif;
    letter-spacing: 0.04em;
    border-radius: 8px;
    font-weight: 500;
}
.stTextInput > div > div > input {
    border-radius: 8px;
    border: 1px solid var(--border);
    font-family: 'DM Sans', sans-serif;
    font-size: 0.88rem;
}
div[data-testid="stExpander"] {
    border: 1px solid var(--border);
    border-radius: 10px;
    background: #FFFFFF;
    box-shadow: var(--shadow);
}
</style>
"""


# ─────────────────────────── QUADRANT LOGIC ─────────────────────────

def assign_quadrant(tc: float, bv: float) -> str:
    if   bv >= 2.0 and tc >= 2.0: return "FAST TRACK"
    elif bv >= 2.0 and tc <  2.0: return "STRATEGIC PRIORITY"
    elif bv <  2.0 and tc >= 2.0: return "OPTIMISE FIRST"
    else:                          return "DEPRIORITISE"


# ─────────────────────────── HELPERS ────────────────────────────────

def load_system_prompt() -> str:
    if not PROMPT_FILE.exists():
        st.error(
            "❌ `scoring_prompt.docx` not found next to `app.py`."
        )
        st.stop()
    doc = Document(str(PROMPT_FILE))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines += [p.text for p in cell.paragraphs]
    return "\n".join(lines)


def extract_pdf_text(uploaded_file) -> str:
    pdf_bytes = uploaded_file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    return "\n\n".join(page.get_text() for page in doc)


def call_claude_scoring(system_prompt: str, procedure_text: str) -> dict:
    client = anthropic.Anthropic()
    msg = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": procedure_text}],
    )
    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$",    "", raw)
    return json.loads(raw)


def call_claude_chat(scoring_data: dict, procedure_text: str,
                     chat_history: list, user_message: str) -> str:
    """Answer follow-up questions about the scorecard."""
    client = anthropic.Anthropic()

    system = f"""You are a Swarovski AI-Readiness Analyst. You have already scored
a procedure document and produced the following scorecard JSON:

{json.dumps(scoring_data, indent=2)}

The original procedure text begins below:
---
{procedure_text[:60000]}
---

Answer the user's questions about the scoring clearly and concisely.
Reference specific evidence from the procedure or scoring rationales.
Keep answers professional, business-focused, and to the point.
Use plain language suitable for senior Swarovski stakeholders.
If asked to re-score or adjust, explain the rationale carefully."""

    messages = []
    for m in chat_history:
        messages.append({"role": m["role"], "content": m["content"]})
    messages.append({"role": "user", "content": user_message})

    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        system=system,
        messages=messages,
    )
    return resp.content[0].text.strip()


def calc_scores(data: dict):
    tc_keys = ["tc1_data_availability","tc2_standardization","tc3_repeatable_process",
               "tc4_systems_integration","tc5_ai_complexity","tc6_delivery_risk",
               "tc7_time_to_implement"]
    bv_keys = ["bv1_financial_value","bv2_sustainability","bv3_market_impact",
               "bv4_customer_experience","bv5_revenue_opportunity","bv6_employee_experience",
               "bv7_compliance_risk"]
    tc_vals = [data.get(k, 0) for k in tc_keys]
    bv_vals = [data.get(k, 0) for k in bv_keys]
    return tc_vals, bv_vals, round(sum(tc_vals)/7, 1), round(sum(bv_vals)/7, 1)


def score_pill(score: int, axis: str) -> str:
    cls = {3: "pill-green", 2: "pill-amber", 1: "pill-red"}.get(score, "pill-blue")
    if axis == "TC":
        lbl = {3: "Low complexity", 2: "Medium", 1: "High complexity"}.get(score, "")
    else:
        lbl = {3: "High value", 2: "Medium", 1: "Low value"}.get(score, "")
    return f'<span class="pill {cls}">{lbl}</span>'


# ─────────────────────────── DOCX BUILDER ───────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def _cell_para(cell, text, bold=False, size=10, color="FFFFFF",
               align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]; p.clear()
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    p.alignment = align

def _cell_big_number(cell, value, bg_hex, font_size=36):
    _set_cell_bg(cell, bg_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_para(cell, value, bold=True, size=font_size,
               color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

def _score_label(score, axis):
    if axis == "TC": return {3:"Low complexity",2:"Medium",1:"High complexity"}.get(score,"")
    return {3:"High value",2:"Medium",1:"Low value"}.get(score,"")

def _quadrant_subtitle(q):
    return {"FAST TRACK":"High Tech Ease · High Business Value",
            "STRATEGIC PRIORITY":"High Business Value · High Complexity",
            "OPTIMISE FIRST":"High Tech Ease · Low Business Value",
            "DEPRIORITISE":"Low Value · High Complexity"}.get(q,"")

def _render_dim_table(doc, dims, data, axis):
    tbl = doc.add_table(rows=len(dims), cols=3)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(0.5)
    tbl.columns[1].width = Inches(2.2)
    tbl.columns[2].width = Inches(4.1)
    for i, (code, label, sk, rk) in enumerate(dims):
        score = data.get(sk, 0); rationale = data.get(rk, "")
        bg = SCORE_BG.get(score, "5D6D7E")
        sc = tbl.rows[i].cells[0]
        sc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(sc, bg)
        _cell_para(sc, str(score), bold=True, size=16, color="FFFFFF",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        lc = tbl.rows[i].cells[1]
        _set_cell_bg(lc, LABEL_BG)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_para(lc, f"{code}  {label}\n\n{_score_label(score,axis)}",
                   bold=True, size=9, color="2C3E8C")
        _cell_para(tbl.rows[i].cells[2], rationale, size=9, color="2C3E50")

def build_scorecard_docx(data: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin=Cm(1.5); section.bottom_margin=Cm(1.5)
        section.left_margin=Cm(2);  section.right_margin=Cm(2)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    tc_vals, bv_vals, tc_dec, bv_dec = calc_scores(data)
    quadrant = assign_quadrant(tc_dec, bv_dec)
    q_emoji  = QUADRANT_EMOJI.get(quadrant, "")
    q_hex    = QUADRANT_HEX.get(quadrant, "5D6D7E")
    exc_count = data.get("exception_clause_count", 0)
    quick_win = data.get("quick_win_match") or None

    # Header
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(5.0); tbl.columns[1].width = Inches(1.8)
    left = tbl.rows[0].cells[0]; _set_cell_bg(left, HEADER_BG)
    p = left.paragraphs[0]; p.clear()
    r = p.add_run("SWAROVSKI · AI-Readiness Scoring Card")
    r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    p2 = left.add_paragraph()
    r2 = p2.add_run(data.get("procedure_title",""))
    r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=RGBColor(0xB0,0xD4,0xFF)
    p3 = left.add_paragraph()
    r3 = p3.add_run(f"Function: {data.get('function','N/A')}")
    r3.font.size=Pt(9); r3.font.color.rgb=RGBColor(0xAA,0xCC,0xFF)
    right = tbl.rows[0].cells[1]; _set_cell_bg(right, q_hex)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_para(right, f"{q_emoji}\n{quadrant}", bold=True, size=10,
               color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Summary — one big number
    sh = doc.add_paragraph("Score Summary")
    sh.runs[0].bold=True; sh.runs[0].font.color.rgb=RGBColor(0x0D,0x1B,0x3E)
    sum_tbl = doc.add_table(rows=2, cols=4)
    sum_tbl.style = "Table Grid"
    for i,(h,v,bg) in enumerate(zip(
        ["Technical Complexity","Business Value","Exception Clauses","Quick-Win Match"],
        [str(tc_dec),str(bv_dec),str(exc_count),"✓" if quick_win else "—"],
        ["2980B9","00C896","8E44AD","16A085"]
    )):
        _set_cell_bg(sum_tbl.rows[0].cells[i], bg)
        _cell_para(sum_tbl.rows[0].cells[i], h, bold=True, size=9, color="FFFFFF",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_big_number(sum_tbl.rows[1].cells[i], v, bg, font_size=36)
    doc.add_paragraph()

    # Metadata
    h2 = doc.add_heading("PROCEDURE METADATA", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x0D,0x1B,0x3E)
    meta_rows = [
        ("Procedure Title", data.get("procedure_title","")),
        ("Function", data.get("function","")),
        ("Process Owner", data.get("process_owner") or "N/A"),
        ("Volume Indicator", data.get("volume_indicator") or "N/A"),
        ("Step Count", str(data.get("step_count",""))),
        ("Systems", ", ".join(data.get("systems_mentioned",[]))),
        ("Exception Clauses", str(exc_count)),
        ("Quick-Win Match", str(quick_win) if quick_win else "None"),
        ("Security Note", data.get("security_flag") or data.get("security_note") or "None flagged"),
    ]
    mt = doc.add_table(rows=len(meta_rows), cols=2)
    mt.style = "Table Grid"
    mt.columns[0].width=Inches(2.0); mt.columns[1].width=Inches(4.8)
    for i,(k,v) in enumerate(meta_rows):
        _set_cell_bg(mt.rows[i].cells[0], LABEL_BG)
        _cell_para(mt.rows[i].cells[0], k, bold=True, size=9, color="0D1B3E")
        _cell_para(mt.rows[i].cells[1], v, size=9, color="2C3E50")
    doc.add_paragraph()

    # TC
    h2 = doc.add_heading("TECHNICAL COMPLEXITY · Axis A  (3=Low/easy · 1=High/hard)", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x29,0x80,0xB9)
    _render_dim_table(doc,[
        ("TC1","Data Availability & Quality","tc1_data_availability","tc1_rationale"),
        ("TC2","Standardization","tc2_standardization","tc2_rationale"),
        ("TC3","Repeatable Process","tc3_repeatable_process","tc3_rationale"),
        ("TC4","Systems & Integration","tc4_systems_integration","tc4_rationale"),
        ("TC5","AI Solution Complexity","tc5_ai_complexity","tc5_rationale"),
        ("TC6","Delivery Risk & Security","tc6_delivery_risk","tc6_rationale"),
        ("TC7","Time to Implement","tc7_time_to_implement","tc7_rationale"),
    ], data, "TC")
    agg = doc.add_paragraph()
    r = agg.add_run(f"TC Aggregate = {tc_dec} / 3.0  (integer: {round(tc_dec)})")
    r.bold=True; r.font.color.rgb=RGBColor(0x29,0x80,0xB9)
    doc.add_paragraph()

    # BV
    h2 = doc.add_heading("BUSINESS VALUE · Axis B  (1=Low · 2=Medium · 3=High)", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x00,0xC8,0x96)
    _render_dim_table(doc,[
        ("BV1","Financial Value","bv1_financial_value","bv1_rationale"),
        ("BV2","Sustainability & Ethical Src.","bv2_sustainability","bv2_rationale"),
        ("BV3","Market Impact","bv3_market_impact","bv3_rationale"),
        ("BV4","Customer Experience","bv4_customer_experience","bv4_rationale"),
        ("BV5","Revenue Opportunity","bv5_revenue_opportunity","bv5_rationale"),
        ("BV6","Employee Experience","bv6_employee_experience","bv6_rationale"),
        ("BV7","Compliance & Risk Reduction","bv7_compliance_risk","bv7_rationale"),
    ], data, "BV")
    agg = doc.add_paragraph()
    r = agg.add_run(f"BV Aggregate = {bv_dec} / 3.0  (integer: {round(bv_dec)})")
    r.bold=True; r.font.color.rgb=RGBColor(0x00,0xC8,0x96)
    doc.add_paragraph()

    # Prioritisation matrix
    h2 = doc.add_heading("PRIORITISATION MATRIX · Quadrant Assignment", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x0D,0x1B,0x3E)
    mx = doc.add_table(rows=1, cols=2)
    mx.style = "Table Grid"
    mx.columns[0].width=Inches(3.5); mx.columns[1].width=Inches(3.3)
    lc = mx.rows[0].cells[0]; _set_cell_bg(lc, LABEL_BG)
    _cell_para(lc,
        "Quadrant Logic  (midpoint = 2.0)\n\n"
        "⚡ FAST TRACK        BV ≥ 2.0  +  TC ≥ 2.0\n"
        "🎯 STRATEGIC         BV ≥ 2.0  +  TC < 2.0\n"
        "🔧 OPTIMISE FIRST    BV < 2.0  +  TC ≥ 2.0\n"
        "🔴 DEPRIORITISE      BV < 2.0  +  TC < 2.0",
        size=9, color="0D1B3E")
    rc = mx.rows[0].cells[1]; _set_cell_bg(rc, q_hex)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_para(rc, f"{q_emoji}  {quadrant}\n\nBV {bv_dec}  ·  TC {tc_dec}\n{_quadrant_subtitle(quadrant)}",
               bold=True, size=12, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Executive summary
    h2 = doc.add_heading("EXECUTIVE SUMMARY", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x0D,0x1B,0x3E)
    ep = doc.add_paragraph(data.get("executive_summary",""))
    ep.runs[0].font.color.rgb = RGBColor(0x2C,0x3E,0x50)
    doc.add_paragraph()

    # JSON
    h2 = doc.add_heading("MACHINE-READABLE OUTPUT · JSON Scoring Card", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x0D,0x1B,0x3E)
    data["quadrant"] = quadrant
    data["tc_aggregate_score_decimal"] = tc_dec
    data["bv_aggregate_score_decimal"] = bv_dec
    jp = doc.add_paragraph(json.dumps(data, indent=2))
    jp.runs[0].font.name="Courier New"; jp.runs[0].font.size=Pt(8)
    jp.runs[0].font.color.rgb=RGBColor(0x34,0x49,0x5E)

    doc.add_paragraph()
    fp = doc.add_paragraph(
        "Swarovski Crystal Business · AI Readiness Assessment · ")
    fp.runs[0].font.size=Pt(8); fp.runs[0].font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ─────────────────────────── STREAMLIT UI ───────────────────────────

def init_state():
    for key, default in {
        "scored_data":     None,
        "procedure_text":  None,
        "chat_messages":   [],
        "docx_bytes":      None,
        "docx_filename":   None,
        "tc_vals":         None,
        "bv_vals":         None,
        "tc_dec":          None,
        "bv_dec":          None,
        "quadrant":        None,
    }.items():
        if key not in st.session_state:
            st.session_state[key] = default


def render_sidebar():
    with st.sidebar:
        st.markdown("""
        <div class="sidebar-logo">
          <span class="diamond">💎</span>
          <span class="brand">SWAROVSKI</span>
          <span class="sub">AI Intelligence Platform</span>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("### Upload Procedure")
        uploaded = st.file_uploader(
            "Drag & drop a PDF here",
            type=["pdf"],
            label_visibility="collapsed",
        )

        # API key
        st.markdown("---")
        st.markdown("### Configuration")
        key = st.secrets["ANTHROPIC_API_KEY"] if "ANTHROPIC_API_KEY" in st.secrets else None
        if not key:
            key = st.text_input("Anthropic API Key", type="password",
                                placeholder="sk-ant-...")
            if key:
                os.environ["ANTHROPIC_API_KEY"] = key

        st.markdown("---")

        # Status indicator
        if st.session_state.scored_data:
            quadrant = st.session_state.quadrant
            q_emoji  = QUADRANT_EMOJI.get(quadrant, "")
            q_hex    = "#" + QUADRANT_HEX.get(quadrant, "5D6D7E")
            st.markdown(f"""
            <div style="background:rgba(255,255,255,0.08);border-radius:10px;
                        padding:1rem;margin-bottom:1rem;border:1px solid rgba(201,168,76,0.25);">
              <div style="font-size:0.68rem;color:rgba(201,168,76,0.8);
                          letter-spacing:0.15em;text-transform:uppercase;margin-bottom:6px;">
                Current Assessment
              </div>
              <div style="font-size:0.85rem;color:#FFFFFF;font-weight:500;margin-bottom:8px;">
                {st.session_state.scored_data.get('procedure_title','')[:40]}
              </div>
              <div style="display:flex;gap:8px;align-items:center;">
                <span style="background:{q_hex};border-radius:6px;padding:3px 10px;
                             font-size:0.72rem;color:#FFFFFF;font-weight:600;">
                  {q_emoji} {quadrant}
                </span>
              </div>
              <div style="margin-top:8px;display:flex;gap:8px;">
                <span style="font-size:0.75rem;color:rgba(255,255,255,0.55);">
                  TC: <strong style="color:#FFFFFF">{st.session_state.tc_dec}</strong>
                </span>
                <span style="font-size:0.75rem;color:rgba(255,255,255,0.55);">
                  BV: <strong style="color:#FFFFFF">{st.session_state.bv_dec}</strong>
                </span>
              </div>
            </div>
            """, unsafe_allow_html=True)

            if st.session_state.chat_messages:
                st.markdown(f"""
                <div style="font-size:0.72rem;color:rgba(255,255,255,0.45);
                            letter-spacing:0.08em;margin-bottom:4px;">
                  CONVERSATION
                </div>
                <div style="font-size:0.8rem;color:rgba(255,255,255,0.65);">
                  {len(st.session_state.chat_messages)//2} question(s) asked
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("""
        <div style="font-size:0.68rem;color:rgba(255,255,255,0.30);
                    letter-spacing:0.06em;line-height:1.6;">
          Swarovski Crystal Business<br>
          AI Readiness Assessment<br>
        </div>
        """, unsafe_allow_html=True)

    return uploaded


def render_welcome():
    st.markdown("""
    <div class="page-header">
      <p class="page-title">AI Readiness Assessment</p>
      <div class="gold-bar"></div>
      <p class="page-sub">
        Analyse procedure documents against Swarovski's AI-readiness framework —
        14 dimensions · Business Value vs. Technical Complexity · Quadrant prioritisation
      </p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([3, 2])
    with col1:
        st.markdown('<div class="hiw-card">', unsafe_allow_html=True)
        st.markdown('<p class="section-title">How It Works</p>', unsafe_allow_html=True)
        steps = [
            ("Upload", "Drag any Swarovski Procedure PDF into the sidebar uploader."),
            ("Score", "AI Assisstant reads the full document and scores 7 Technical Complexity and 7 Business Value dimensions."),
            ("Quadrant", "The procedure is placed in one of four delivery zones: ⚡ Fast Track · 🎯 Strategic · 🔧 Optimise · 🔴 Deprioritise."),
            ("Download", "A formatted scorecard (.docx) is generated matching the standard Swarovski template."),
            ("Converse", "Ask follow-up questions about any score, rationale, or recommendation directly in the chat panel."),
        ]
        for i,(title,desc) in enumerate(steps,1):
            st.markdown(f"""
            <div class="step-item">
              <div class="step-num">{i}</div>
              <div class="step-text"><strong>{title}</strong> — {desc}</div>
            </div>
            """, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div style="background:linear-gradient(135deg,#0D1B3E 0%,#1E3469 100%);
                    border-radius:14px;padding:2rem;height:100%;box-sizing:border-box;">
          <p style="font-family:'Cormorant Garamond',serif;font-size:1.8rem;
                    color:#FFFFFF;font-weight:700;margin:0 0 1rem;line-height:1.2;">
            Prioritisation<br>Matrix
          </p>
          <div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;margin-top:1rem;">
            <div style="background:#00C896;border-radius:10px;padding:1rem;text-align:center;">
              <div style="font-size:1.4rem;">⚡</div>
              <div style="font-size:0.7rem;color:#FFFFFF;font-weight:600;margin-top:4px;">FAST TRACK</div>
              <div style="font-size:0.6rem;color:rgba(255,255,255,0.75);margin-top:2px;">BV≥2 · TC≥2</div>
            </div>
            <div style="background:#3D8FD1;border-radius:10px;padding:1rem;text-align:center;">
              <div style="font-size:1.4rem;">🎯</div>
              <div style="font-size:0.7rem;color:#FFFFFF;font-weight:600;margin-top:4px;">STRATEGIC</div>
              <div style="font-size:0.6rem;color:rgba(255,255,255,0.75);margin-top:2px;">BV≥2 · TC&lt;2</div>
            </div>
            <div style="background:#F0A500;border-radius:10px;padding:1rem;text-align:center;">
              <div style="font-size:1.4rem;">🔧</div>
              <div style="font-size:0.7rem;color:#FFFFFF;font-weight:600;margin-top:4px;">OPTIMISE FIRST</div>
              <div style="font-size:0.6rem;color:rgba(255,255,255,0.75);margin-top:2px;">BV&lt;2 · TC≥2</div>
            </div>
            <div style="background:#E05555;border-radius:10px;padding:1rem;text-align:center;">
              <div style="font-size:1.4rem;">🔴</div>
              <div style="font-size:0.7rem;color:#FFFFFF;font-weight:600;margin-top:4px;">DEPRIORITISE</div>
              <div style="font-size:0.6rem;color:rgba(255,255,255,0.75);margin-top:2px;">BV&lt;2 · TC&lt;2</div>
            </div>
          </div>
        </div>
        """, unsafe_allow_html=True)


def render_results():
    data      = st.session_state.scored_data
    tc_vals   = st.session_state.tc_vals
    bv_vals   = st.session_state.bv_vals
    tc_dec    = st.session_state.tc_dec
    bv_dec    = st.session_state.bv_dec
    quadrant  = st.session_state.quadrant
    q_emoji   = QUADRANT_EMOJI.get(quadrant, "")
    q_hex     = QUADRANT_HEX.get(quadrant, "5D6D7E")

    # Page header
    st.markdown(f"""
    <div class="page-header">
      <p class="page-title">{data.get('procedure_title','Assessment Results')}</p>
      <div class="gold-bar"></div>
      <p class="page-sub">
        Function: {data.get('function','—')} &nbsp;·&nbsp;
        Process Owner: {data.get('process_owner','—')} &nbsp;·&nbsp;
        Systems: {', '.join(data.get('systems_mentioned',[]))}
      </p>
    </div>
    """, unsafe_allow_html=True)

    # Three big score cards
    st.markdown(f"""
    <div class="score-grid">
      <div class="score-card" style="background:linear-gradient(135deg,#2980B9,#1A5F8A);">
        <div class="label">Technical Complexity</div>
        <div class="value">{tc_dec}</div>
        <div class="unit">out of 3.0 &nbsp;·&nbsp; 3 = easiest to automate</div>
      </div>
      <div class="score-card" style="background:linear-gradient(135deg,#00B87A,#007A5E);">
        <div class="label">Business Value</div>
        <div class="value">{bv_dec}</div>
        <div class="unit">out of 3.0 &nbsp;·&nbsp; 3 = highest impact</div>
      </div>
      <div class="score-card" style="background:linear-gradient(135deg,#{q_hex},{''.join(hex(max(0,int(q_hex[i:i+2],16)-30)).replace('0x','').zfill(2) for i in (0,2,4))});">
        <div class="label">Quadrant Assignment</div>
        <div class="value" style="font-size:2.6rem;">{q_emoji}</div>
        <div class="unit" style="font-size:0.8rem;font-weight:600;">{quadrant}</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Quadrant banner
    st.markdown(f"""
    <div class="quadrant-banner" style="background:linear-gradient(135deg,#{q_hex} 0%,{''.join(hex(max(0,int(q_hex[i:i+2],16)-40)).replace('0x','').zfill(2) for i in (0,2,4))} 100%);">
      <span class="q-emoji">{q_emoji}</span>
      <div>
        <div class="q-title">{quadrant}</div>
        <div class="q-sub">{QUADRANT_LABEL.get(quadrant,'')}</div>
      </div>
      <div class="q-scores">
        <span>Business Value</span>
        <strong>{bv_dec}</strong>
        <span style="margin-top:6px;">Technical Complexity</span>
        <strong>{tc_dec}</strong>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Executive summary
    st.markdown('<p class="section-title">Executive Summary</p>', unsafe_allow_html=True)
    st.markdown(f'<div class="exec-card">{data.get("executive_summary","")}</div>',
                unsafe_allow_html=True)

    # Download strip — always visible
    safe = re.sub(r"[^\w\s-]","",data.get("procedure_title","scorecard"))[:60]
    fname = f"{safe.replace(' ','_')}_Scorecard.docx"
    st.markdown('<p class="section-title">Scorecard Download</p>', unsafe_allow_html=True)
    st.download_button(
        label="⬇️  Download Scorecard (.docx)",
        data=st.session_state.docx_bytes,
        file_name=fname,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        type="primary",
    )

    # Dimension details (collapsible)
    with st.expander("📊 Technical Complexity — All 7 Dimensions", expanded=False):
        tc_dims = [
            ("TC1","Data Availability & Quality","tc1_data_availability","tc1_rationale"),
            ("TC2","Standardization","tc2_standardization","tc2_rationale"),
            ("TC3","Repeatable Process","tc3_repeatable_process","tc3_rationale"),
            ("TC4","Systems & Integration","tc4_systems_integration","tc4_rationale"),
            ("TC5","AI Solution Complexity","tc5_ai_complexity","tc5_rationale"),
            ("TC6","Delivery Risk & Security","tc6_delivery_risk","tc6_rationale"),
            ("TC7","Time to Implement","tc7_time_to_implement","tc7_rationale"),
        ]
        for (code,label,sk,rk), score in zip(tc_dims, tc_vals):
            bg  = {"3":"#00C896","2":"#F0A500","1":"#E05555"}.get(str(score),"#8899AA")
            rationale = data.get(rk,"")
            st.markdown(f"""
            <div class="dim-row">
              <div class="dim-badge" style="background:{bg};">{score}</div>
              <div class="dim-label">
                <span class="code">{code}</span>
                <span class="name">{label}</span>
                <span class="tag">{_score_label(score,'TC')}</span>
              </div>
              <div class="dim-rationale">{rationale}</div>
            </div>
            """, unsafe_allow_html=True)

    with st.expander("💼 Business Value — All 7 Dimensions", expanded=False):
        bv_dims = [
            ("BV1","Financial Value","bv1_financial_value","bv1_rationale"),
            ("BV2","Sustainability & Ethical Sourcing","bv2_sustainability","bv2_rationale"),
            ("BV3","Market Impact","bv3_market_impact","bv3_rationale"),
            ("BV4","Customer Experience","bv4_customer_experience","bv4_rationale"),
            ("BV5","Revenue Opportunity","bv5_revenue_opportunity","bv5_rationale"),
            ("BV6","Employee Experience","bv6_employee_experience","bv6_rationale"),
            ("BV7","Compliance & Risk Reduction","bv7_compliance_risk","bv7_rationale"),
        ]
        for (code,label,sk,rk), score in zip(bv_dims, bv_vals):
            bg  = {"3":"#00C896","2":"#F0A500","1":"#E05555"}.get(str(score),"#8899AA")
            rationale = data.get(rk,"")
            st.markdown(f"""
            <div class="dim-row">
              <div class="dim-badge" style="background:{bg};">{score}</div>
              <div class="dim-label">
                <span class="code">{code}</span>
                <span class="name">{label}</span>
                <span class="tag">{_score_label(score,'BV')}</span>
              </div>
              <div class="dim-rationale">{rationale}</div>
            </div>
            """, unsafe_allow_html=True)

    with st.expander("📋 Procedure Metadata", expanded=False):
        exc_count = data.get("exception_clause_count", 0)
        quick_win = data.get("quick_win_match") or "None"
        rows = [
            ("Procedure Title",   data.get("procedure_title","")),
            ("Function",          data.get("function","")),
            ("Process Owner",     data.get("process_owner") or "N/A"),
            ("Volume Indicator",  data.get("volume_indicator") or "N/A"),
            ("Step Count",        str(data.get("step_count",""))),
            ("Systems Involved",  ", ".join(data.get("systems_mentioned",[]))),
            ("Exception Clauses", str(exc_count)),
            ("Quick-Win Match",   str(quick_win)),
            ("Security Note",     data.get("security_flag") or
                                  data.get("security_note") or "None flagged"),
        ]
        rows_html = "".join(
            f"<tr><td>{k}</td><td>{v}</td></tr>" for k,v in rows
        )
        st.markdown(
            f'<table class="meta-table">{rows_html}</table>',
            unsafe_allow_html=True
        )


def render_chat():
    st.markdown('<p class="section-title">Ask the AI Analyst</p>', unsafe_allow_html=True)

    # Chat header
    st.markdown("""
    <div class="chat-container">
      <div class="chat-header">
        <span class="chat-dot"></span>
        <div>
          <div class="ch-title">AI Analyst · Live Session</div>
          <div class="ch-sub">Ask any question about the scores, rationales or recommendations</div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Render chat history
    if not st.session_state.chat_messages:
        st.markdown("""
        <div style="padding:1.2rem;background:#FFFFFF;border:1px solid rgba(13,27,62,0.10);
                    border-radius:10px;margin-top:8px;">
          <div class="msg-bot" style="max-width:100%;margin:0;">
            <span class="msg-avatar-bot">💎</span>
            Hello. I have completed the AI-readiness assessment for this procedure.
            Feel free to ask me anything — why a particular dimension was scored the way it was,
            what the quadrant means for prioritisation, or how the process compares to
            Swarovski's quick-win use cases.
          </div>
        </div>
        """, unsafe_allow_html=True)
    else:
        msgs_html = '<div style="padding:1rem 1.2rem;background:#FFFFFF;border:1px solid rgba(13,27,62,0.10);border-radius:10px;margin-top:8px;max-height:480px;overflow-y:auto;">'
        msgs_html += '<div class="msg-bot"><span class="msg-avatar-bot">💎</span>Hello. Assessment complete. Ask me anything about the scores.</div>'
        for m in st.session_state.chat_messages:
            if m["role"] == "user":
                msgs_html += f'<div class="msg-user">{m["content"]}</div>'
            else:
                msgs_html += f'<div class="msg-bot"><span class="msg-avatar-bot">💎</span>{m["content"]}</div>'
        msgs_html += "</div>"
        st.markdown(msgs_html, unsafe_allow_html=True)

    # Suggested questions
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("""
    <p style="font-size:0.72rem;letter-spacing:0.1em;text-transform:uppercase;
              color:#4A5980;margin-bottom:8px;">Suggested Questions</p>
    """, unsafe_allow_html=True)

    suggestions = [
        "Why is the Technical Complexity score low?",
        "What would improve the Business Value score?",
        "What does the quadrant assignment mean for our roadmap?",
        "Which dimension has the most room for improvement?",
    ]
    cols = st.columns(2)
    for i, suggestion in enumerate(suggestions):
        if cols[i % 2].button(suggestion, key=f"sugg_{i}", use_container_width=True):
            _handle_chat_input(suggestion)
            st.rerun()

    # Chat input
    st.markdown("<br>", unsafe_allow_html=True)
    with st.form("chat_form", clear_on_submit=True):
        col_input, col_send = st.columns([5, 1])
        user_input = col_input.text_input(
            "Your question",
            placeholder="e.g. Why did TC2 Standardization score a 3?",
            label_visibility="collapsed",
        )
        submitted = col_send.form_submit_button("Send →", use_container_width=True)

    if submitted and user_input.strip():
        _handle_chat_input(user_input.strip())
        st.rerun()


def _handle_chat_input(user_message: str):
    st.session_state.chat_messages.append({"role":"user","content":user_message})
    with st.spinner("Analysing…"):
        try:
            reply = call_claude_chat(
                st.session_state.scored_data,
                st.session_state.procedure_text,
                st.session_state.chat_messages[:-1],
                user_message,
            )
        except Exception as e:
            reply = f"Sorry, I encountered an error: {e}"
    st.session_state.chat_messages.append({"role":"assistant","content":reply})


# ─────────────────────────── MAIN ───────────────────────────────────

def main():
    st.markdown(CSS, unsafe_allow_html=True)
    init_state()
    uploaded = render_sidebar()

    # ── Handle new upload + generate ──
    if uploaded is not None and os.environ.get("ANTHROPIC_API_KEY"):
        # Only re-score if this is a newly uploaded file
        if (st.session_state.scored_data is None or
                uploaded.name != st.session_state.get("_last_filename")):

            st.session_state["_last_filename"] = uploaded.name

            col_btn, _ = st.columns([2, 3])
            if col_btn.button("🚀 Generate Scorecard", type="primary", use_container_width=True):
                with st.spinner("📖 Extracting PDF text…"):
                    try:
                        proc_text = extract_pdf_text(uploaded)
                    except Exception as e:
                        st.error(f"PDF extraction failed: {e}"); return

                with st.spinner("🤖 AI Assistant is scoring the procedure across 14 dimensions…"):
                    try:
                        sys_prompt = load_system_prompt()
                        data = call_claude_scoring(sys_prompt, proc_text)
                    except json.JSONDecodeError as e:
                        st.error(f"AI Assistant returned non-JSON output: {e}"); return
                    except Exception as e:
                        st.error(f"API error: {e}"); return

                tc_vals, bv_vals, tc_dec, bv_dec = calc_scores(data)
                quadrant = assign_quadrant(tc_dec, bv_dec)

                with st.spinner("📝 Building scorecard .docx…"):
                    docx_bytes = build_scorecard_docx(data)

                safe  = re.sub(r"[^\w\s-]","",data.get("procedure_title","scorecard"))[:60]
                fname = f"{safe.replace(' ','_')}_Scorecard.docx"

                # Persist everything
                st.session_state.scored_data    = data
                st.session_state.procedure_text = proc_text
                st.session_state.tc_vals        = tc_vals
                st.session_state.bv_vals        = bv_vals
                st.session_state.tc_dec         = tc_dec
                st.session_state.bv_dec         = bv_dec
                st.session_state.quadrant       = quadrant
                st.session_state.docx_bytes     = docx_bytes
                st.session_state.docx_filename  = fname
                st.session_state.chat_messages  = []   # reset chat for new doc
                st.rerun()
            else:
                if st.session_state.scored_data is None:
                    render_welcome()
            return

    # ── Show results (persistent) ──
    if st.session_state.scored_data is not None:
        render_results()
        st.markdown("<br>", unsafe_allow_html=True)
        render_chat()
    else:
        render_welcome()
        if uploaded and not os.environ.get("ANTHROPIC_API_KEY"):
            st.warning("Please enter your Anthropic API key in the sidebar to continue.")


if __name__ == "__main__":
    main()
